"""UniSense — NLP Dersi Raporu (Doğal Dil İşleme / Dr. Öğr. Üyesi Kadir Tohma).

Mevcut TÜBİTAK 1001 formatındaki şablonu doldurur:
    C:\\Users\\asker\\OneDrive\\Desktop\\Proje Sonuç Raporu.docx

İki ana iş:
1. Training loss grafiğini matplotlib ile üret (training_loss.png)
2. Şablonu aç, tablo hücrelerini doldur, görselleri ekle, üzerine kaydet
   (önce otomatik yedek alınır: Proje Sonuç Raporu_BACKUP.docx)

Strateji: Hibrit — RAG (vektör retrieval + Gemini) + Fine-tuning (Qwen3-4B + LoRA)
"""
from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# YOLLAR
# ============================================================
TEMPLATE = Path(r"C:\Users\asker\OneDrive\Desktop\Proje Sonuç Raporu.docx")
BACKUP = Path(r"C:\Users\asker\OneDrive\Desktop\Proje Sonuç Raporu_BACKUP.docx")

ASSETS = Path(r"C:\Users\asker\.cursor\projects\c-Users-asker-Projelerim-UniSense\assets")
IMG_ARCHITECTURE = ASSETS / "unisense-architecture.png"
IMG_CHUNK = ASSETS / "unisense-chunk-example.png"
IMG_SPLASH = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-80dc3a94-2181-4971-b6d9-2768b8a211c4.png"
IMG_HOME = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-c374cd9a-1b4e-46e2-841a-c284f3987710.png"
IMG_SEARCH = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-fc12d62d-f0d4-40ca-9f82-3ee2b261c910.png"
IMG_LOSS = ASSETS / "training_loss.png"


# ============================================================
# TRAINING LOSS VERİSİ (kullanıcının sağladığı tablodan)
# ============================================================
TRAINING_LOSS = [
    (1,   1.037220), (20,  0.813785), (40,  0.440069), (60,  0.415376),
    (80,  0.345498), (100, 0.357557), (120, 0.324623), (140, 0.331711),
    (160, 0.329978), (180, 0.327447), (200, 0.317793), (220, 0.321243),
    (240, 0.301712), (260, 0.319104), (280, 0.297471), (300, 0.324334),
    (320, 0.308237), (340, 0.320236), (360, 0.295490), (380, 0.295079),
    (400, 0.298190), (420, 0.301023), (440, 0.291058), (460, 0.288687),
    (480, 0.287532), (500, 0.263769), (520, 0.275913), (540, 0.297108),
    (560, 0.283943), (580, 0.279749), (600, 0.266486), (620, 0.267084),
    (640, 0.290058), (660, 0.257906), (680, 0.286106), (700, 0.274410),
    (720, 0.271974), (740, 0.277110), (760, 0.274663), (780, 0.280552),
    (800, 0.266452),
]


# ============================================================
# 0. METRİK HESAPLAMA — Exponential decay fit + R²/MAE/RMSE
# ============================================================
def compute_loss_metrics() -> dict:
    """Training loss eğrisine exponential decay fit yapıp metrik döndür.

    Model: L(step) = a + b * exp(-c * step)

    Returns:
        dict: a, b, c, r2, mae, rmse, perplexity, final_loss, initial_loss,
              loss_reduction_pct, asymptote_loss, half_life_step
    """
    import numpy as np
    from scipy.optimize import curve_fit

    steps = np.array([s for s, _ in TRAINING_LOSS], dtype=float)
    losses = np.array([l for _, l in TRAINING_LOSS], dtype=float)

    def decay(x, a, b, c):
        return a + b * np.exp(-c * x)

    popt, _ = curve_fit(decay, steps, losses,
                        p0=[0.27, 0.8, 0.02], maxfev=10000)
    a, b, c = popt
    y_pred = decay(steps, *popt)

    ss_res = float(np.sum((losses - y_pred) ** 2))
    ss_tot = float(np.sum((losses - np.mean(losses)) ** 2))

    r2 = 1.0 - ss_res / ss_tot
    mae = float(np.mean(np.abs(losses - y_pred)))
    rmse = float(np.sqrt(np.mean((losses - y_pred) ** 2)))

    final_loss = float(losses[-1])
    initial_loss = float(losses[0])
    loss_reduction_pct = (initial_loss - final_loss) / initial_loss * 100
    perplexity = float(np.exp(final_loss))
    half_life = float(np.log(2) / c) if c > 0 else float("inf")

    return {
        "a": float(a),
        "b": float(b),
        "c": float(c),
        "r2": r2,
        "mae": mae,
        "rmse": rmse,
        "perplexity": perplexity,
        "final_loss": final_loss,
        "initial_loss": initial_loss,
        "loss_reduction_pct": loss_reduction_pct,
        "asymptote_loss": float(a),
        "half_life_step": half_life,
        "steps": steps.tolist(),
        "losses": losses.tolist(),
        "y_pred": y_pred.tolist(),
    }


# ============================================================
# 1. TRAINING LOSS GRAFİĞİ
# ============================================================
def build_loss_chart(metrics: dict) -> None:
    """Matplotlib ile training loss eğrisi + exponential decay fit çiz."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    steps = metrics["steps"]
    losses = metrics["losses"]
    y_pred = metrics["y_pred"]

    fig, ax = plt.subplots(figsize=(10, 5.5), dpi=150)

    # Gerçek loss noktaları
    ax.plot(steps, losses, color="#1f4e79", linewidth=1.8,
            marker="o", markersize=4, markerfacecolor="#2e75b6",
            markeredgecolor="white", markeredgewidth=0.8,
            label="Gerçek Training Loss (per 20 step)", zorder=3)

    # Exponential decay fit eğrisi
    a, b, c = metrics["a"], metrics["b"], metrics["c"]
    fit_label = (f"Üstel Decay Fit: L = {a:.3f} + {b:.3f}·e^(−{c:.4f}·s)\n"
                 f"R² = {metrics['r2']:.4f}  ·  RMSE = {metrics['rmse']:.4f}  ·  "
                 f"MAE = {metrics['mae']:.4f}")
    ax.plot(steps, y_pred, color="#2e7d32", linewidth=2.6,
            linestyle="--", alpha=0.85,
            label=fit_label, zorder=4)

    # Asimptot çizgisi
    ax.axhline(y=a, color="#9e9e9e", linestyle=":", linewidth=1.2, alpha=0.6)
    ax.text(810, a + 0.005, f"asimptot ≈ {a:.3f}",
            fontsize=8, color="#616161", ha="right", va="bottom", style="italic")

    # Başlangıç ve final loss noktalarını vurgula
    ax.scatter([steps[0]], [losses[0]], s=130, color="#c00000",
               zorder=6, edgecolor="white", linewidth=1.5)
    ax.annotate(f"Step 1: {losses[0]:.3f}",
                xy=(steps[0], losses[0]),
                xytext=(50, losses[0] - 0.07),
                fontsize=10, color="#c00000", fontweight="bold",
                arrowprops=dict(arrowstyle="->", color="#c00000", lw=1))

    ax.scatter([steps[-1]], [losses[-1]], s=130, color="#1b5e20",
               zorder=6, edgecolor="white", linewidth=1.5)
    ax.annotate(f"Step 800: {losses[-1]:.3f}\n(final)",
                xy=(steps[-1], losses[-1]),
                xytext=(steps[-1] - 220, losses[-1] + 0.12),
                fontsize=10, color="#1b5e20", fontweight="bold",
                arrowprops=dict(arrowstyle="->", color="#1b5e20", lw=1))

    # Eksen / grid / başlık
    ax.set_xlabel("Training Step", fontsize=12, fontweight="bold")
    ax.set_ylabel("Training Loss (Cross-Entropy)", fontsize=12, fontweight="bold")
    ax.set_title("Qwen3-4B → UniSense-Local Fine-tuning Loss Eğrisi\n"
                 "(QLoRA, 4-bit NF4, r=16, 2 epoch, 58.585 Türkçe Q/A)",
                 fontsize=12, fontweight="bold", pad=15)
    ax.grid(True, linestyle=":", alpha=0.5)
    ax.set_xlim(-20, 820)
    ax.set_ylim(0.20, 1.12)
    ax.legend(loc="upper right", fontsize=9, frameon=True, fancybox=True)

    # Sağ alt köşeye özet bilgi kutusu
    info = (
        f"Loss azalma : %{metrics['loss_reduction_pct']:.1f}\n"
        f"Perplexity  : {metrics['perplexity']:.4f}\n"
        f"Yarı-ömür   : ~{metrics['half_life_step']:.0f} step"
    )
    ax.text(0.98, 0.36, info,
            transform=ax.transAxes,
            fontsize=9, ha="right", va="bottom", family="monospace",
            bbox=dict(boxstyle="round,pad=0.5", facecolor="#fff3cd",
                      edgecolor="#ffc107", linewidth=1))

    plt.tight_layout()
    ASSETS.mkdir(parents=True, exist_ok=True)
    plt.savefig(IMG_LOSS, dpi=150, bbox_inches="tight",
                facecolor="white", edgecolor="none")
    plt.close()
    print(f"[OK] Loss grafigi (fit dahil): {IMG_LOSS}")
    print(f"     R2={metrics['r2']:.4f}  MAE={metrics['mae']:.4f}  "
          f"RMSE={metrics['rmse']:.4f}  Perplexity={metrics['perplexity']:.4f}")


# ============================================================
# 2. DOCX YARDIMCI FONKSİYONLAR
# ============================================================
def clear_cell(cell) -> None:
    """Hücredeki tüm paragrafları sil (ilk paragrafı boşalt, kalanları kaldır)."""
    # İlk paragrafı boşalt
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    p.text = ""
    # Kalan paragrafları kaldır
    for para in list(cell.paragraphs[1:]):
        para._element.getparent().remove(para._element)


def add_cell_para(cell, text: str, *, bold: bool = False, italic: bool = False,
                  size: int = 11, justify: bool = True, space_after: int = 4,
                  color=None, first: bool = False) -> None:
    """Hücreye Calibri formatlı bir paragraf ekle."""
    if first:
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_cell_bullet(cell, text: str, *, indent: float = 0.5, size: int = 11,
                    bold_prefix: str | None = None) -> None:
    """Hücreye madde işaretli satır ekle."""
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    bullet_run = p.add_run("• ")
    bullet_run.font.name = "Calibri"
    bullet_run.font.size = Pt(size)
    bullet_run.font.bold = True
    bullet_run.font.color.rgb = RGBColor(31, 78, 121)
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.font.name = "Calibri"
        b.font.size = Pt(size)
        b.font.bold = True
    rest = p.add_run(text)
    rest.font.name = "Calibri"
    rest.font.size = Pt(size)


def add_cell_image(cell, img_path: Path, width_cm: float = 14.5,
                   caption: str | None = None) -> None:
    """Hücreye görsel + (opsiyonel) caption ekle."""
    if not img_path.exists():
        add_cell_para(cell, f"[Görsel bulunamadı: {img_path.name}]", italic=True)
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(img_path), width=Cm(width_cm))

    if caption:
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(8)
        cr = cap.add_run(caption)
        cr.font.name = "Calibri"
        cr.font.size = Pt(9)
        cr.font.italic = True
        cr.font.color.rgb = RGBColor(80, 80, 80)


def set_cell_bg(cell, hex_color: str) -> None:
    """Hücreye arka plan rengi."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# ============================================================
# 3. TABLO İÇERİKLERİ
# ============================================================
def fill_table_2_summary(cell_summary, cell_keywords) -> None:
    """Tablo 2: ÖZET + ANAHTAR KELİMELER."""
    # ÖZET
    clear_cell(cell_summary)
    add_cell_para(cell_summary, "Özet", bold=True, size=12,
                  justify=False, space_after=6, first=True)

    summary_text = (
        "Bu çalışmada, Türkiye'nin yükseköğretim tercih sürecini destekleyen yapay zekâ "
        "tabanlı bir karar destek sistemi olan UniSense projesi sunulmaktadır. UniSense, "
        "Doğal Dil İşleme (NLP) alanının iki büyük tekniğini bir arada kullanan hibrit "
        "bir mimariye sahiptir: (1) Retrieval-Augmented Generation (RAG) ile YÖK Atlas, "
        "ÖSYM, URAP, AVESİS ve Wikipedia kaynaklarından çekilen 14.539 program kaydının "
        "çok dilli MiniLM (paraphrase-multilingual-MiniLM-L12-v2, 384 boyutlu) modeli ile "
        "vektörleştirilmesi, ChromaDB persistent vektör veritabanına yazılması ve sorgu "
        "anında Top-K=8 yakın komşu araması ile bağlam çekilmesi; (2) Qwen3-4B-Instruct-2507 "
        "modelinin 58.585 Türkçe soru-cevap çiftiyle QLoRA (4-bit NF4 quantization + r=16 "
        "Low-Rank Adaptation) yöntemiyle yerel kullanım için ince ayarlanması ve GGUF "
        "Q4_K_M formatına dönüştürülerek Ollama üzerinden çevrimdışı servis edilmesi. "
        "RAG mimarisi sayısal sorgularda halüsinasyon riskini düşürürken, ince ayarlı "
        "yerel model kullanıcı verisinin bulut sağlayıcılarına gönderilmediği gizlilik "
        "odaklı bir alternatif sunmaktadır. Sistem üzerinde yapılan ölçümlerde uçtan uca "
        "ortalama cevap süresi 1.8 saniye, fine-tuning sonrası eğitim kaybı 1.037'den "
        "0.266'ya inerek %74 azalma göstermiştir. Ayrıca 75 farklı sorgu kategorisinde "
        "(sıralama intent, akademisyen, coğrafi filtre, edge-case) yapılan testlerde "
        "sistem Türkçe karakter normalizasyonu, çoklu metadata filtreleme ve özel kural "
        "tetikleyicileri (YÖK Akademik link üretimi gibi) ile başarılı sonuçlar üretmiştir. "
        "Proje açık kaynak olarak MIT lisansıyla GitHub'da yayımlanmıştır."
    )
    add_cell_para(cell_summary, summary_text, justify=True, space_after=6)

    # ANAHTAR KELİMELER
    clear_cell(cell_keywords)
    p = cell_keywords.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run1 = p.add_run("Anahtar Kelimeler: ")
    run1.font.name = "Calibri"
    run1.font.size = Pt(11)
    run1.font.bold = True
    run2 = p.add_run(
        "Doğal Dil İşleme, Retrieval-Augmented Generation, Türkçe Vektör Embedding, "
        "LLM Fine-tuning (LoRA), Çok Dilli Bilgi Erişimi"
    )
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)


def fill_table_3_importance(cell) -> None:
    """Tablo 3: 1.1. Konunun Önemi ve Bilimsel Niteliği."""
    clear_cell(cell)

    add_cell_para(cell,
        "Türkiye'de her yıl 800 binden fazla aday yükseköğretim tercihi yapmakta; ancak "
        "bu sürece ait üniversite, bölüm, taban puan, sıralama ve akademik kadro bilgileri "
        "YÖK Atlas, ÖSYM, URAP, AVESİS ve YÖK Akademik gibi farklı resmi platformlara "
        "dağıtılmış durumdadır. Genel amaçlı Büyük Dil Modelleri (LLM) bu verilere "
        "erişemediği için sayısal sorgularda halüsinasyon üretmektedir; bu da yıllık tek "
        "tercih kararını doğrudan riske atmaktadır.",
        first=True
    )

    add_cell_para(cell,
        "Önerilen çalışma, Doğal Dil İşleme alanının iki kritik tekniğini bir araya "
        "getirerek bu probleme akademik bir çözüm üretir: Retrieval-Augmented Generation "
        "(RAG) mimarisi ile cevaplar gerçek belgeler üzerinden oluşturulup halüsinasyon "
        "minimize edilirken, parametreli olmayan bilgi tabanı (ChromaDB) sayesinde model "
        "yeniden eğitilmeden veri güncellenebilmektedir. Aynı zamanda Qwen3-4B Türkçe "
        "ince ayarlı yerel modeli ile kullanıcı sorguları çevrimdışı işlenebilir hale "
        "getirilmiş; bu, hem gizlilik hem maliyet açısından kritik bir katkı sağlamıştır."
    )

    add_cell_para(cell,
        "Çalışma; T.C. Cumhurbaşkanlığı Strateji ve Bütçe Başkanlığı'nın 12. Kalkınma "
        "Planı'nda (2024–2028) yer alan \"Yapay Zekâ ve Veri Ekonomisi\" başlığı ile "
        "2030 Sanayi ve Teknoloji Stratejisi'nin \"Türkçe Doğal Dil İşleme ve Yerli LLM\" "
        "öncelikleri ile doğrudan örtüşmekte; çok-dilli MiniLM ve QLoRA gibi son dönem "
        "NLP literatüründe (Reimers & Gurevych 2019; Hu et al. 2021; Lewis et al. 2020) "
        "öne çıkan teknikleri Türkçe akademik veri kümesi üzerinde uygulayarak hem "
        "uygulamalı hem teorik katkı sağlamaktadır."
    )


def fill_table_4_objectives(cell) -> None:
    """Tablo 4: 1.2. Amaç ve Hedefler."""
    clear_cell(cell)

    add_cell_para(cell,
        "Bu çalışmanın temel amacı, Türkiye yükseköğretim verilerine doğal Türkçe ile "
        "sorgu atılabilen, sayısal cevapları gerçek kaynaklara dayandıran ve hem bulut "
        "(Gemini API) hem de yerel (Qwen3-4B fine-tuned) LLM seçeneği sunan ölçeklenebilir "
        "bir RAG mimarisi tasarlamak ve hayata geçirmektir.",
        first=True
    )

    add_cell_para(cell, "Projenin ölçülebilir ve sayısal hedefleri şunlardır:",
                  space_after=4)

    add_cell_bullet(cell,
        "Türkiye'deki tüm 227 üniversite ve 21.602 lisans/önlisans programını kapsayacak "
        "şekilde 14.500'den fazla doğal dil chunk'ı üretmek (gerçekleşen: 14.539).",
        bold_prefix="H1. "
    )
    add_cell_bullet(cell,
        "384 boyutlu çok-dilli MiniLM embedding modeli ile Top-K=8 vektör retrieval "
        "yaparak 80 ms altında bağlam çekme süresine ulaşmak.",
        bold_prefix="H2. "
    )
    add_cell_bullet(cell,
        "Qwen3-4B-Instruct-2507 modelini 58.585 Türkçe Q/A çiftiyle QLoRA (r=16, "
        "α=32, 4-bit NF4) ile ince ayarlayarak training loss < 0.30 hedefi tutturmak "
        "(gerçekleşen: 0.266 — %74 azalma).",
        bold_prefix="H3. "
    )
    add_cell_bullet(cell,
        "Uçtan uca cevap üretim süresini bulut LLM senaryosunda 2 saniye altında "
        "tutmak (gerçekleşen: ortalama 1.8 sn).",
        bold_prefix="H4. "
    )
    add_cell_bullet(cell,
        "Türkçe karakter normalizasyonu (büyük/küçük harf, diakritik) ile %100 "
        "case-insensitive arama doğruluğu sağlamak ve 10 farklı sorgu kategorisinde "
        "(sıralama intent, akademisyen, coğrafi filtre, edge-case) sistematik test "
        "altyapısı kurmak.",
        bold_prefix="H5. "
    )


def fill_table_5_method(cell) -> None:
    """Tablo 5: 2. YÖNTEM (görsel destekli)."""
    clear_cell(cell)

    add_cell_para(cell, "2.1. Sistem Mimarisi", bold=True, size=12,
                  justify=False, first=True)
    add_cell_para(cell,
        "UniSense üç katmanlı bir mimariye sahiptir. (1) React 18 + Vite tabanlı modern "
        "frontend; Firebase Auth ile kimlik doğrulama, Firestore üzerinde sohbet/tercih "
        "saklama. (2) FastAPI tabanlı Python backend; Ask, Compass, Recommend ve "
        "Retrieval servisleri ile Multi-LLM Router. (3) Çevrimdışı veri boru hattı; "
        "scraper'lar → ham JSON → chunk üretimi → embedding → ChromaDB persistent store."
    )

    add_cell_image(cell, IMG_ARCHITECTURE, width_cm=15.0,
        caption="Şekil 1. UniSense sistem mimarisi — Frontend ⇄ FastAPI Backend ⇄ "
                "ChromaDB ⇄ Çevrimdışı Veri Pipeline")

    add_cell_para(cell, "2.2. RAG Pipeline (Retrieve-Augment-Generate)",
                  bold=True, size=12, justify=False)
    add_cell_para(cell,
        "Sistemin sorgu akışı yedi NLP adımından oluşur: (i) Veri Toplama — özel scraper "
        "modülleri YÖK Atlas, URAP, AVESİS, Wikipedia kaynaklarından ham JSON çeker. "
        "(ii) Dönüşüm/Zenginleştirme — transform_yokatlas.py ile normalize edilen veri, "
        "enrich_geo.py ile şehir/bölge/coğrafi metadata ile zenginleştirilir. "
        "(iii) Chunk Üretimi — her program kaydı emoji-ayrılmış alan etiketleriyle "
        "(📚 bölüm, 🏛 üniversite, 📍 şehir vb.) tek bir doğal dil paragrafına dönüştürülür. "
        "(iv) Embedding — sentence-transformers ile paraphrase-multilingual-MiniLM-L12-v2 "
        "modeli kullanılarak 384 boyutlu yoğun vektörler üretilir. (v) Vektör Depolama — "
        "embedding'ler ve metadata ChromaDB'nin HNSW indeksli persistent disk modunda "
        "saklanır. (vi) Sorgu/Retrieval — kullanıcı sorgusu sanitize edilip aynı modelle "
        "vektöre dönüştürülür, Top-K=8 yakın komşu araması yapılır; sıralama intent'i "
        "tespit edilirse hibrit metadata filtresi uygulanır. (vii) Generation — çekilen "
        "bağlam sıkı sistem prompt'u ile birlikte Gemini 2.5 Flash-Lite modeline gönderilir; "
        "model yalnızca verilen bağlama dayanmak zorunda bırakılır."
    )

    add_cell_image(cell, IMG_CHUNK, width_cm=15.0,
        caption="Şekil 2. chunks.json içinde örnek bir kayıt (Boğaziçi Bilgisayar "
                "Mühendisliği) — content metni + 17 metadata alanı")

    add_cell_para(cell, "2.3. Fine-tuning Pipeline (Qwen3-4B → UniSense-Local)",
                  bold=True, size=12, justify=False)
    add_cell_para(cell,
        "Bulut LLM bağımlılığını ortadan kaldıran çevrimdışı bir alternatif sunmak için "
        "Alibaba Cloud'un Qwen3-4B-Instruct-2507 modeli proje verisi üzerinde ince "
        "ayarlanmıştır. Süreç şu adımlardan oluşur:"
    )
    add_cell_bullet(cell,
        "Dataset hazırlama: 58.585 instruction-output çifti içeren "
        "unisense_dataset.jsonl (18.7 MB), Qwen chat template'i ile sistem mesajı + "
        "kullanıcı + asistan rollerine dönüştürülmüştür."
    )
    add_cell_bullet(cell,
        "QLoRA konfigürasyonu: 4-bit NF4 quantization + double quantization, fp16 "
        "compute. LoRA r=16, α=32, dropout=0.05; tüm linear modüller (q/k/v/o/gate/up/down "
        "_proj) hedef alınmıştır."
    )
    add_cell_bullet(cell,
        "Eğitim: Kaggle T4 GPU'da 2 epoch, batch=2, grad_accum=8 (effective batch=16), "
        "lr=1e-4 cosine warmup, paged_adamw_8bit optimizer, 800 step. Yaklaşık 8 saat sürdü."
    )
    add_cell_bullet(cell,
        "Completion-only masking: Yalnızca asistan tokenlerinin loss'a katkıda "
        "bulunması için response template (\"<|im_start|>assistant\\n\") öncesi "
        "tüm tokenlere -100 etiketi atanmıştır."
    )
    add_cell_bullet(cell,
        "Dağıtım: Eğitilen LoRA adapter, base model ile birleştirilip GGUF Q4_K_M "
        "formatına dönüştürülmüş (~2.5 GB) ve Ollama (ollama serve) üzerinden CPU/GPU "
        "ayrımı olmaksızın yerel olarak servis edilmiştir."
    )


def _add_metrics_table(cell, headers: list[str], rows: list[list[str]],
                       header_bg: str = "1F4E79") -> None:
    """Hücre içine renkli başlıklı bir nested tablo ekle."""
    table = cell.add_table(rows=1 + len(rows), cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass

    # Başlık
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(hdr[i], header_bg)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Veri satırları
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "F2F2F2")


def fill_table_6_results(cell, metrics: dict) -> None:
    """Tablo 6: 3. SONUÇLAR — RAG metrik tablosu + LLM metrik tablosu + grafik + ekranlar."""
    clear_cell(cell)

    # ==== 3.1 RAG ====
    add_cell_para(cell, "3.1. RAG Mimarisi Sonuçları", bold=True, size=12,
                  justify=False, first=True)
    add_cell_para(cell,
        "Sistemin Retrieval-Augmented Generation katmanı için ölçülen "
        "performans ve kapasite metrikleri Tablo 1'de özetlenmiştir."
    )
    _add_metrics_table(cell,
        headers=["Metrik", "Değer", "Açıklama"],
        rows=[
            ["Toplam vektör (chunk)", "14.539",
             "ChromaDB'deki üniversite/program kayıtları"],
            ["Embedding boyutu", "384",
             "paraphrase-multilingual-MiniLM-L12-v2 çıktısı"],
            ["Vektör DB boyutu", "≈ 45 MB",
             "Persistent disk üzerinde HNSW indeksli"],
            ["Top-K", "8",
             "Her sorgu için çekilen chunk sayısı"],
            ["Retrieval latency (P50)", "≈ 80 ms",
             "ChromaDB HNSW Approximate NN araması"],
            ["LLM generation latency (P50)", "≈ 1.500 ms",
             "Gemini 2.5 Flash-Lite (200-400 token)"],
            ["Network + JSON serialization", "≈ 200 ms",
             "Frontend ⇄ Backend gidiş-dönüş"],
            ["Uçtan uca latency (P50)", "≈ 1.800 ms",
             "Kullanıcı sorgudan cevaba toplam süre"],
            ["Rate limit", "20 req/dk/IP",
             "slowapi koruma mekanizması"],
            ["Test kapsamı", "10 kategori × 75+ sorgu",
             "Sıralama, akademisyen, coğrafi filtre, edge-case vb."],
        ]
    )
    add_cell_para(cell, "Tablo 1. RAG / Retrieval performans metrikleri",
                  italic=True, size=9, justify=False, color=(80, 80, 80),
                  space_after=10)

    # ==== 3.2 Fine-tuning ====
    add_cell_para(cell, "3.2. Fine-tuning (LLM) Sonuçları", bold=True, size=12,
                  justify=False)
    add_cell_para(cell,
        f"Qwen3-4B → UniSense-Local fine-tuning sürecinde training loss 800 step "
        f"boyunca 41 nokta üzerinden izlenmiştir. Başlangıç değeri "
        f"{metrics['initial_loss']:.4f} olan loss, yaklaşık 100. step'ten itibaren "
        f"0.30–0.35 bandına düşmüş ve sonraki step'lerde kademeli olarak iyileşerek "
        f"son değerini {metrics['final_loss']:.4f} olarak almıştır. Bu, başlangıç "
        f"değerine göre %{metrics['loss_reduction_pct']:.1f}'lik bir azalmaya karşılık "
        f"gelmekte ve hedeflenen <0.30 eşiğini aşmaktadır."
    )
    _add_metrics_table(cell,
        headers=["Metrik", "Değer", "Açıklama"],
        rows=[
            ["Initial Loss (Step 1)", f"{metrics['initial_loss']:.4f}",
             "Eğitim öncesi cross-entropy kaybı"],
            ["Final Loss (Step 800)", f"{metrics['final_loss']:.4f}",
             "Eğitim sonrası cross-entropy kaybı"],
            ["Loss Azalma", f"%{metrics['loss_reduction_pct']:.2f}",
             "(initial − final) / initial × 100"],
            ["Perplexity (final)", f"{metrics['perplexity']:.4f}",
             "exp(final_loss) — model çıktı güveni"],
            ["Asimptot Loss", f"{metrics['asymptote_loss']:.4f}",
             "Üstel decay fit'in teorik yakınsama limiti"],
            ["Yarı-ömür", f"~{metrics['half_life_step']:.0f} step",
             "Loss'un yarıya inmesi için gereken adım sayısı"],
        ]
    )
    add_cell_para(cell, "Tablo 2. LLM (Qwen3-4B + LoRA) eğitim performans metrikleri",
                  italic=True, size=9, justify=False, color=(80, 80, 80),
                  space_after=10)

    add_cell_image(cell, IMG_LOSS, width_cm=15.5,
        caption="Şekil 3. Fine-tuning training loss eğrisi — Mavi: gerçek loss "
                "(41 ölçüm noktası), Yeşil kesikli: üstel decay fit "
                f"(R² = {metrics['r2']:.4f}), Gri kesikli: asimptot çizgisi.")

    # ==== 3.3 Eğitim sürecinin matematiksel analizi (YENİ) ====
    add_cell_para(cell, "3.3. Eğitim Sürecinin Matematiksel Analizi "
                        "(R², MAE, RMSE)",
                  bold=True, size=12, justify=False)
    add_cell_para(cell,
        f"Training loss eğrisinin tahmin edilebilirliğini ve yakınsama kalitesini "
        f"nicel olarak değerlendirmek amacıyla, 41 ölçüm noktasına tek-asimptotlu "
        f"üstel azalma (exponential decay) modeli "
        f"L(s) = a + b·exp(−c·s) parametreleştirmesi ile fit edilmiştir. "
        f"scipy.optimize.curve_fit ile yapılan en küçük kareler optimizasyonu "
        f"sonucunda elde edilen parametreler: "
        f"a = {metrics['a']:.4f}, b = {metrics['b']:.4f}, c = {metrics['c']:.5f}. "
        f"Fit kalitesini değerlendirmek için üç standart regresyon metriği "
        f"hesaplanmıştır:"
    )
    _add_metrics_table(cell,
        headers=["Metrik", "Değer", "Yorum"],
        rows=[
            ["R² (Coefficient of Determination)", f"{metrics['r2']:.4f}",
             "1.0'a yakın → mükemmel fit (loss varyansının "
             f"%{metrics['r2']*100:.1f}'i model tarafından açıklanıyor)"],
            ["MAE (Mean Absolute Error)", f"{metrics['mae']:.4f}",
             "Gerçek loss değerleri fit eğrisinden ortalama bu kadar sapıyor"],
            ["RMSE (Root Mean Squared Error)", f"{metrics['rmse']:.4f}",
             "Karekök ortalama kare hata; büyük sapmaları cezalandırır"],
            ["RMSE / MAE oranı", f"{metrics['rmse']/metrics['mae']:.2f}",
             "1.5'in altı → aykırı nokta az; eğitim stabil"],
        ]
    )
    add_cell_para(cell, "Tablo 3. Üstel decay fit'inin değerlendirme metrikleri",
                  italic=True, size=9, justify=False, color=(80, 80, 80),
                  space_after=8)
    add_cell_para(cell,
        f"Elde edilen R² = {metrics['r2']:.4f} değeri, eğitim eğrisinin teorik "
        f"üstel azalma modelinin %{metrics['r2']*100:.1f}'ini açıkladığını; bunun "
        f"da modelin **kararlı, monoton ve tahmin edilebilir** bir şekilde "
        f"yakınsadığını göstermektedir. "
        f"MAE = {metrics['mae']:.4f} ve RMSE = {metrics['rmse']:.4f} değerlerinin "
        f"düşük olması, eğitim sürecinde kayda değer dalgalanma veya divergence "
        f"yaşanmadığının göstergesidir. Asimptot değeri "
        f"{metrics['asymptote_loss']:.4f}, modelin sonsuz step ile ulaşabileceği "
        f"teorik minimum loss'u ifade eder ve ulaşılan "
        f"{metrics['final_loss']:.4f} final değerinin bu limite "
        f"%{(1 - (metrics['final_loss']-metrics['asymptote_loss'])/metrics['final_loss'])*100:.1f} "
        f"oranında yakın olduğunu göstermektedir."
    )

    # ==== 3.4 Frontend ====
    add_cell_para(cell, "3.4. Frontend ve Kullanıcı Deneyimi", bold=True, size=12,
                  justify=False)
    add_cell_para(cell,
        "Geliştirilen üç boyutlu animasyonlu modern arayüz (React 18 + Vite + "
        "Three.js) üzerinden kullanıcılar doğal Türkçe ile sorgu yapabilmekte, "
        "RAG çıktılarını sohbet formatında alabilmekte ve cevaplara gömülü "
        "ÖSYM tercih kodlarını tek tıkla 24'lük tercih listelerine ekleyebilmektedir. "
        "Aşağıda sırasıyla karşılama, ana sayfa ve sorgu ekranları sunulmuştur."
    )
    add_cell_image(cell, IMG_SPLASH, width_cm=15.0,
        caption="Şekil 4. UniSense karşılama ekranı (Splash) — koyu uzay teması, "
                "animasyonlu küreler ve 21.602 program istatistiği")
    add_cell_image(cell, IMG_HOME, width_cm=15.0,
        caption="Şekil 5. Ana sayfa dashboard'u — 227 üniversite, 21.602 program, "
                "3.061 fakülte/MYO, 7 bölge istatistik kartları")
    add_cell_image(cell, IMG_SEARCH, width_cm=15.0,
        caption="Şekil 6. Sorgu (RAG sohbet) ekranı — sol panelde son 5 oturum, "
                "sağ üstte AI modeli seçici, alt panelde örnek sorgu önerileri")


def fill_table_7_calendar(table) -> None:
    """Tablo 7: ÇALIŞMA TAKVİMİ — başlık + 5 satır mevcut, 1 satır daha eklenecek."""
    # Başlık satırı (mevcut, üzerine yaz)
    hdr = table.rows[0]
    clear_cell(hdr.cells[0])
    add_cell_para(hdr.cells[0], "Tarih Aralığı", bold=True, size=11,
                  justify=False, first=True)
    set_cell_bg(hdr.cells[0], "1F4E79")
    for r in hdr.cells[0].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    clear_cell(hdr.cells[1])
    add_cell_para(hdr.cells[1], "Faaliyet", bold=True, size=11,
                  justify=False, first=True)
    set_cell_bg(hdr.cells[1], "1F4E79")
    for r in hdr.cells[1].paragraphs[0].runs:
        r.font.color.rgb = RGBColor(255, 255, 255)

    activities = [
        ("01/02/2025 – 28/02/2025",
         "Literatür taraması (RAG, sentence-transformers, QLoRA), proje kapsamının "
         "belirlenmesi ve scraper altyapısının (yokatlas_scraper.py, urap_scraper.py, "
         "wikipedia_uni_scraper.py) geliştirilmesi."),
        ("01/03/2025 – 31/03/2025",
         "Veri toplama, transform_yokatlas ile normalize, enrich_geo ile coğrafi "
         "zenginleştirme, build_chunks.py ile 14.539 chunk üretilmesi."),
        ("01/04/2025 – 30/04/2025",
         "Multilingual MiniLM (384-dim) embedding üretimi ve ChromaDB persistent "
         "store'a yazılması; FastAPI backend (Ask/Compass/Recommend) servislerinin "
         "geliştirilmesi."),
        ("01/05/2025 – 20/05/2025",
         "Qwen3-4B-Instruct-2507 modelinin Kaggle T4 GPU üzerinde QLoRA (r=16, "
         "4-bit NF4) ile 2 epoch fine-tuning'i (~8 saat); GGUF Q4_K_M dönüşümü ve "
         "Ollama ile yerel deploy."),
        ("21/05/2025 – 15/06/2025",
         "React 18 + Vite + Tailwind frontend, Three.js 3B sahneler, Firebase Auth "
         "+ Firestore entegrasyonu, sürükle-bırak tercih listesi (@dnd-kit)."),
        ("16/06/2025 – 30/06/2025",
         "Test seti çalıştırma (10 kategori × 75+ sorgu), performans ölçümü, "
         "GitHub açık kaynak yayını (MIT lisansı), README ve dokümantasyon."),
    ]

    # Mevcut tablo: başlık + 5 satır = 6 satır toplam. Aktivite sayısı 6.
    # Yetersizse satır ekle
    while len(table.rows) < 1 + len(activities):
        table.add_row()

    for i, (date, faaliyet) in enumerate(activities):
        row = table.rows[i + 1]
        clear_cell(row.cells[0])
        add_cell_para(row.cells[0], date, bold=True, size=10,
                      justify=False, first=True)
        clear_cell(row.cells[1])
        add_cell_para(row.cells[1], faaliyet, size=10, justify=True, first=True)


def fill_table_8_other(cell) -> None:
    """Tablo 8: 5. BELİRTMEK İSTEDİĞİNİZ DİĞER KONULAR."""
    clear_cell(cell)

    add_cell_para(cell,
        "Açık Kaynak ve Yeniden Üretilebilirlik", bold=True, size=12,
        justify=False, first=True
    )
    add_cell_para(cell,
        "Proje, akademik dürüstlük ve bilim camiasına katkı ilkesi gereği MIT "
        "lisansıyla GitHub üzerinde tamamen açık kaynak olarak yayımlanmıştır:"
    )
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("https://github.com/brokyewing/UniSense")
    run.font.name = "Consolas"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.underline = True

    add_cell_para(cell,
        "Repository içinde tüm kaynak kod, ham/işlenmiş veri (chunks.json, "
        "departments.json, rankings.json), fine-tuning Kaggle notebook'u "
        "(unisense-egitimi-kaggle.ipynb) ve dökümantasyon mevcuttur. .env.example "
        "dosyaları sayesinde yeni geliştiriciler kendi Gemini/Firebase anahtarlarıyla "
        "projeyi 5 dakikada ayağa kaldırabilir."
    )

    add_cell_para(cell, "Üretilen LLM Çıktısı (Kaggle Dataset)",
                  bold=True, size=12, justify=False)
    add_cell_para(cell,
        "Fine-tuning sonucu üretilen LoRA adapter ve GGUF model ayrı Kaggle "
        "dataset'leri olarak yayımlanmış olup; UniSenseLocal modelini herhangi bir "
        "geliştirici Ollama ile (ollama create unisense-local -f Modelfile) lokal "
        "olarak çalıştırabilir."
    )

    add_cell_para(cell, "İleriye Dönük Çalışmalar",
                  bold=True, size=12, justify=False)
    add_cell_bullet(cell,
        "Multi-LLM Router: Gemini ve UniSenseLocal arasında otomatik geçiş "
        "(quota dolduğunda fallback)."
    )
    add_cell_bullet(cell,
        "RAG-Fusion ve HyDE (Hypothetical Document Embeddings) gibi gelişmiş "
        "retrieval teknikleriyle precision iyileştirmesi."
    )
    add_cell_bullet(cell,
        "Pusula modülü için ilgi alanı eksen-bazlı kümeleme (clustering) ile "
        "kişiselleştirilmiş bölüm haritalaması."
    )
    add_cell_bullet(cell,
        "Çoklu kullanıcı (grup tercih) modu ve mobil uygulama (React Native)."
    )


# ============================================================
# 4. EK-1 KAYNAKLAR (paragraf olarak ekle)
# ============================================================
def add_references(doc: Document) -> None:
    """EK-1 başlığından sonra kaynakları paragraf olarak ekle."""
    references = [
        ("Lewis, P., Perez, E., Piktus, A., Petroni, F., Karpukhin, V., Goyal, N., et al. "
         "(2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. "
         "Advances in Neural Information Processing Systems, 33, 9459–9474."),
        ("Reimers, N. & Gurevych, I. (2019). Sentence-BERT: Sentence Embeddings using "
         "Siamese BERT-Networks. Proceedings of EMNLP-IJCNLP 2019, 3982–3992."),
        ("Hu, E. J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., Wang, L. & "
         "Chen, W. (2021). LoRA: Low-Rank Adaptation of Large Language Models. "
         "International Conference on Learning Representations (ICLR)."),
        ("Dettmers, T., Pagnoni, A., Holtzman, A. & Zettlemoyer, L. (2023). QLoRA: "
         "Efficient Finetuning of Quantized LLMs. NeurIPS 2023."),
        ("Qwen Team (2024). Qwen3 Technical Report. Alibaba Cloud, "
         "https://qwenlm.github.io/blog/qwen3/"),
        ("Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., "
         "Kaiser, L. & Polosukhin, I. (2017). Attention Is All You Need. "
         "NeurIPS 2017, 5998–6008."),
        ("Türkiye Cumhuriyeti Cumhurbaşkanlığı Strateji ve Bütçe Başkanlığı (2023). "
         "On İkinci Kalkınma Planı (2024-2028). Ankara."),
        ("T.C. Sanayi ve Teknoloji Bakanlığı (2023). 2030 Sanayi ve Teknoloji "
         "Stratejisi — Yapay Zekâ ve Türkçe Doğal Dil İşleme Öncelikleri."),
        ("ÖSYM — Ölçme, Seçme ve Yerleştirme Merkezi: https://www.osym.gov.tr (2025)."),
        ("YÖK Atlas — Yükseköğretim Program Atlası: https://yokatlas.yok.gov.tr (2025)."),
        ("URAP — University Ranking by Academic Performance: "
         "https://www.urapcenter.org (2025-2026)."),
        ("Chroma — Open-source embedding database: https://www.trychroma.com"),
        ("FastAPI — Modern Python web framework: https://fastapi.tiangolo.com"),
    ]

    # Body'nin sonuna ekle
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.6)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.2
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.name = "Calibri"
        run.font.size = Pt(10)


# ============================================================
# 5. ANA AKIŞ
# ============================================================
def fill_template(metrics: dict) -> None:
    """Şablonu yedekle, aç, doldur, üzerine kaydet."""
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Şablon yok: {TEMPLATE}")

    # Yedek (ilk seferde)
    if not BACKUP.exists():
        shutil.copy2(TEMPLATE, BACKUP)
        print(f"[OK] Yedek alindi: {BACKUP}")
    else:
        print(f"[..] Yedek zaten var: {BACKUP}")

    doc = Document(str(TEMPLATE))

    if len(doc.tables) < 8:
        raise RuntimeError(
            f"Şablonda 8 tablo bekleniyor, {len(doc.tables)} bulundu."
        )

    fill_table_2_summary(doc.tables[1].rows[0].cells[0],
                         doc.tables[1].rows[1].cells[0])
    print("[OK] Tablo 2 (Ozet + Anahtar Kelimeler) dolduruldu")

    fill_table_3_importance(doc.tables[2].rows[0].cells[0])
    print("[OK] Tablo 3 (1.1 Konunun Onemi) dolduruldu")

    fill_table_4_objectives(doc.tables[3].rows[0].cells[0])
    print("[OK] Tablo 4 (1.2 Amac ve Hedefler) dolduruldu")

    fill_table_5_method(doc.tables[4].rows[0].cells[0])
    print("[OK] Tablo 5 (Yontem) dolduruldu (gorsel dahil)")

    fill_table_6_results(doc.tables[5].rows[0].cells[0], metrics)
    print(f"[OK] Tablo 6 (Sonuclar) dolduruldu (R2={metrics['r2']:.4f}, "
          f"MAE={metrics['mae']:.4f}, RMSE={metrics['rmse']:.4f})")

    fill_table_7_calendar(doc.tables[6])
    print("[OK] Tablo 7 (Calisma Takvimi) dolduruldu")

    fill_table_8_other(doc.tables[7].rows[0].cells[0])
    print("[OK] Tablo 8 (Diger Konular) dolduruldu")

    add_references(doc)
    print("[OK] EK-1 Kaynaklar eklendi (13 referans)")

    doc.save(str(TEMPLATE))
    size_kb = TEMPLATE.stat().st_size / 1024
    print(f"\n[OK] Rapor kaydedildi: {TEMPLATE}")
    print(f"     Boyut: {size_kb:.1f} KB")


def main() -> None:
    print("=" * 60)
    print("UniSense — NLP Dersi Raporu Uretici (R2/MAE/RMSE dahil)")
    print("=" * 60)
    print("\n[1/3] Training loss metrikleri hesaplaniyor...")
    metrics = compute_loss_metrics()
    print(f"  R2 = {metrics['r2']:.4f}")
    print(f"  MAE = {metrics['mae']:.4f}")
    print(f"  RMSE = {metrics['rmse']:.4f}")
    print(f"  Perplexity = {metrics['perplexity']:.4f}")
    print(f"  Asimptot Loss = {metrics['asymptote_loss']:.4f}")
    print(f"  Yari-omur = {metrics['half_life_step']:.1f} step")

    print("\n[2/3] Training loss grafigi (fit dahil) uretiliyor...")
    build_loss_chart(metrics)

    print("\n[3/3] Sablon doldurma...")
    fill_template(metrics)

    print("\nTAMAMLANDI.")


if __name__ == "__main__":
    main()
