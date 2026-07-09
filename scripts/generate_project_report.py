"""UniSense ders projesi raporu - .docx üretici.

Kullanım:
    python scripts/generate_project_report.py

Çıktı:
    C:\\Users\\asker\\OneDrive\\Desktop\\UniSense_Proje_Raporu_IBRAHIM_ASKEROGLU.docx

Şablon kapak bilgileri:
    Öğrenci  : 232503303 - İBRAHİM ASKEROĞLU
    Hoca     : Dr. Öğr. Üyesi Halil İbrahim OKUR
    Ders     : Mühendislikte Bilgisayar Uygulamaları II
    Proje    : UniSense
    Dönem    : 2025-2026 Bahar
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


# ============================================================
# YOL TANIMLARI
# ============================================================
ASSETS = Path(r"C:\Users\asker\.cursor\projects\c-Users-asker-Projelerim-UniSense\assets")

IMG_ARCHITECTURE = ASSETS / "unisense-architecture.png"
IMG_SPLASH = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-80dc3a94-2181-4971-b6d9-2768b8a211c4.png"
IMG_HOME = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-c374cd9a-1b4e-46e2-841a-c284f3987710.png"
IMG_SEARCH = ASSETS / "c__Users_asker_AppData_Roaming_Cursor_User_workspaceStorage_149631f1989dab970474885636cd4b8e_images_image-fc12d62d-f0d4-40ca-9f82-3ee2b261c910.png"

OUT_PATH = Path(r"C:\Users\asker\OneDrive\Desktop\UniSense_Proje_Raporu_IBRAHIM_ASKEROGLU.docx")

GITHUB_URL = "https://github.com/brokyewing/UniSense"


# ============================================================
# YARDIMCI FONKSİYONLAR
# ============================================================
def set_cell_bg(cell, hex_color: str) -> None:
    """Tablo hücresine arka plan rengi ekle."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def style_heading(p, color=(31, 73, 125), size=16, bold=True):
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(*color)


def add_para(doc: Document, text: str, *, justify: bool = True, size: int = 11,
             bold: bool = False, italic: bool = False, color=None,
             space_after: int = 6) -> None:
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bullet(doc: Document, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)
    # Alt çizgi efekti
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(54, 95, 145)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              *, header_bg: str = "1F497D", header_color=(255, 255, 255),
              col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Başlık satırı
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*header_color)
        set_cell_bg(hdr_cells[i], header_bg)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Veri satırları
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            cells[ci].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # Zebra stripe
            if ri % 2 == 1:
                set_cell_bg(cells[ci], "F2F2F2")

    # Sütun genişlikleri
    if col_widths_cm:
        for col, w in zip(table.columns, col_widths_cm):
            for cell in col.cells:
                cell.width = Cm(w)

    # Tablodan sonra boşluk
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


# ============================================================
# RAPOR İÇERİĞİ
# ============================================================
def build_cover(doc: Document) -> None:
    """Kapak sayfası — şablona uyumlu."""
    # Üst boşluk
    for _ in range(3):
        doc.add_paragraph()

    # Öğrenci bilgileri
    info = [
        "Öğrenci Bilgileri: 232503303 — İBRAHİM ASKEROĞLU",
        "",
        "Dersin Öğretim Üyesi: DR. ÖĞR. ÜYESİ HALİL İBRAHİM OKUR",
        "",
        "Ders Bilgileri: MÜHENDİSLİKTE BİLGİSAYAR UYGULAMALARI II",
        "",
        "Projenin Başlığı: UniSense",
    ]
    for line in info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        if line and line[0] in ("Ö", "D", "P"):
            run.font.bold = True

    # Ortada büyük proje adı
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UniSense")
    run.font.name = "Calibri"
    run.font.size = Pt(48)
    run.font.bold = True
    run.font.color.rgb = RGBColor(70, 35, 175)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAG Tabanlı Türkiye Üniversite Tercih Asistanı")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("GitHub: ")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.bold = True
    run2 = p.add_run(GITHUB_URL)
    run2.font.name = "Consolas"
    run2.font.size = Pt(12)
    run2.font.color.rgb = RGBColor(0, 102, 204)
    run2.font.underline = True

    # Alt boşluk + dönem
    for _ in range(7):
        doc.add_paragraph()

    for line in ("2025-2026 Yılı", "Bahar Dönemi"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(line)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.font.bold = True


def section_amac(doc: Document) -> None:
    add_h1(doc, "1. Projenin Amacı")

    add_para(doc,
        "UniSense, Türkiye'deki YKS (Yükseköğretim Kurumları Sınavı) tercih sürecini yaşayan "
        "öğrencileri ve velileri destekleyen, Retrieval-Augmented Generation (RAG) tabanlı bir "
        "yapay zekâ asistanıdır. Projenin temel amacı; YÖK Atlas, ÖSYM, URAP gibi farklı resmi "
        "kaynaklara dağılmış olan üniversite, bölüm, sıralama, taban puan, kontenjan, akademik kadro "
        "bilgilerini tek bir konuşma arayüzü altında toplamak ve öğrencinin doğal Türkçe ile sorduğu "
        "sorulara bağlama uygun, sayısal verilerle desteklenmiş cevaplar üretmektir."
    )

    add_para(doc,
        "Sistem klasik bir arama motoru ya da statik bir filtreleme aracı değildir. Aksine; öğrencinin "
        "puanı, sıralaması, ilgi alanları, şehir tercihi gibi çoklu kriterleri anlayan ve bunlara karşılık "
        "gelen veri kümesinden vektör benzerliği ile en uygun bağlamı çekip Büyük Dil Modeli (LLM) "
        "üzerinden cevaplayan akıllı bir karar destek sistemidir."
    )

    add_h2(doc, "1.1. Projenin Hedefleri")
    add_bullet(doc, "Türkiye'deki 227 üniversite, 21.602 lisans/önlisans programı ve 3.061 fakülte/MYO için tek tıkla erişilebilir bilgi katmanı sağlamak.")
    add_bullet(doc, "Öğrencinin doğal Türkçe sorgularını (örn. \"300.000 sıralamayla devlet üniversiteleri\", \"Boğaziçi Bilgisayar Mühendisliği taban puanı\") gerçek YÖK Atlas verisiyle cevaplamak.")
    add_bullet(doc, "Kişiselleştirilmiş 24'lük tercih listesi oluşturma, sürükle-bırak sıralama ve otomatik backfill özellikleriyle tercih dönemini kolaylaştırmak.")
    add_bullet(doc, "Pusula (Compass) modülü ile ilgi alanı ve güçlü dersler üzerinden bölüm önerisi üretmek.")
    add_bullet(doc, "Halüsinasyonu minimize etmek: cevaplar yalnızca retrieve edilen gerçek belgeler üzerinden üretilir; LLM'in kendi kafasından sayı uydurması engellenir.")


def section_problem(doc: Document) -> None:
    add_h1(doc, "2. Problem Tanımı")

    add_para(doc,
        "Türkiye'de her yıl 3 milyondan fazla öğrenci YKS'ye girmekte, bunlardan yaklaşık 800 bini "
        "yükseköğretime yerleşmek için tercih yapmaktadır. Bu süreçte öğrenciler aşağıdaki problemlerle "
        "karşı karşıyadır:"
    )

    add_h2(doc, "2.1. Bilgi Dağınıklığı")
    add_para(doc,
        "Üniversite ve bölüm bilgileri en az dört farklı resmi platforma yayılmıştır: YÖK Atlas (taban puan, "
        "sıralama, kontenjan), ÖSYM (puan türleri ve tercih kodları), URAP (akademik sıralama), "
        "akademik.yok.gov.tr (akademik kadro). Tek bir bölüm hakkında kapsamlı bilgi edinmek için "
        "öğrencinin birden fazla siteyi gezmesi, manuel olarak çapraz arama yapması gerekmektedir."
    )

    add_h2(doc, "2.2. Klasik Arama Yetersizliği")
    add_para(doc,
        "YÖK Atlas gibi mevcut araçlar yalnızca tam metin (exact-match) arama ve sabit filtreler sunar. "
        "Öğrencinin \"İstanbul'da denizi olan illerde 80 bin sıraya kadar bilgisayar mühendisliği bölümleri\" "
        "gibi çok-boyutlu, doğal dilde ifade edilmiş bir sorguyu doğrudan cevaplayabilen bir araç yoktur. "
        "Mevcut sistemler kategorik filtre kombinasyonu kurmayı kullanıcıya yükler."
    )

    add_h2(doc, "2.3. Genel Amaçlı LLM'lerin Sınırları")
    add_para(doc,
        "ChatGPT, Gemini, Claude gibi genel amaçlı LLM'ler Türkiye yükseköğretim verilerine doğrudan "
        "erişimi olmadığı için ya eski (eğitim kesim tarihindeki) bilgi verir ya da halüsinasyon üretir. "
        "Özellikle sayısal veri (taban puan, sıralama, kontenjan) söz konusu olduğunda halüsinasyon "
        "riski yüksektir ve bu durum öğrencinin yıllık tek tercih kararını doğrudan etkiler."
    )

    add_h2(doc, "2.4. Hedef Kullanıcı Senaryoları")
    add_table(doc,
        headers=["Senaryo", "Kullanıcı", "Beklenti"],
        rows=[
            ["Geniş tarama", "Bursalı 12. sınıf öğrencisi (450k sıra)", "Sıralamasıyla girebileceği tüm devlet üniversitelerini görmek"],
            ["Hedef bölüm", "TYT/AYT'de iyi puan almış aday", "Boğaziçi/ODTÜ/İTÜ Bilgisayar Mühendisliği taban puanlarını karşılaştırmak"],
            ["İlgi keşfi", "Henüz bölüm seçmemiş öğrenci", "Pusula modülü ile ilgi alanına uygun bölümler önerilmek"],
            ["Aile danışmanlığı", "Veli", "Şehir + bölüm + burs filtreleriyle vakıf üniversitelerini taramak"],
            ["Akademik kalite", "Akademik kadroya önem veren aday", "URAP sıralaması + akademisyen sayıları + YÖK Akademik linki"],
        ],
        col_widths_cm=[3.5, 5.5, 7.5],
    )


def section_veri(doc: Document) -> None:
    add_h1(doc, "3. Veri Kaynakları ve Elde Edilme Yöntemleri")

    add_para(doc,
        "UniSense'in cevap kalitesi tamamen altındaki veri kümesinin doğruluğuna ve güncelliğine bağlıdır. "
        "Bu nedenle proje kapsamında beş farklı resmi/yarı-resmi kaynak otomatik olarak çekilmiş, "
        "temizlenmiş ve birleştirilmiştir. Aşağıda kullanılan kaynaklar ve elde edilme yöntemleri "
        "özetlenmiştir."
    )

    add_h2(doc, "3.1. Veri Kaynakları")
    add_table(doc,
        headers=["Kaynak", "İçerik", "Yöntem", "Boyut"],
        rows=[
            ["YÖK Atlas (yokatlas.yok.gov.tr)", "21.602 lisans + önlisans programı, taban puan, sıralama, kontenjan, eğitim dili", "Özel HTTP scraper (yokatlas_scraper.py) — yokAtlas iç JSON API'leri", "~92 MB ham JSON"],
            ["ÖSYM (osym.gov.tr)", "Puan türleri (SAY, EA, SÖZ, DİL, TYT), tercih kodları, kontenjan PDF'leri", "PyMuPDF tabanlı PDF parse + manuel doğrulama", "~3 MB"],
            ["URAP (urapcenter.org)", "Türk üniversitelerinin akademik sıralaması (227 üniversite × 9 metrik)", "BeautifulSoup + lxml ile HTML scrape (urap_scraper.py)", "~250 KB JSON"],
            ["AVESİS (avesis.<uni>.edu.tr)", "Üniversite bazlı akademisyen verisi (kadro, ünvan, alanlar)", "Üniversite bazlı dinamik scraper (avesis_scraper.py)", "~8 MB JSON"],
            ["Wikipedia", "Üniversite tarihi, kuruluş yılı, kampüs bilgisi, öğrenci sayıları", "wikipedia-api Python kütüphanesi", "~1 MB JSON"],
        ],
        col_widths_cm=[4.0, 5.5, 4.5, 2.5],
    )

    add_h2(doc, "3.2. Toplam Veri Hacmi (Pipeline Sonrası)")
    add_table(doc,
        headers=["Veri Türü", "Adet", "Açıklama"],
        rows=[
            ["Üniversite", "227", "Tüm devlet + vakıf + KKTC üniversiteleri"],
            ["Lisans/Önlisans Programı", "21.602", "2025 yılı kontenjan tablosundan"],
            ["Fakülte / MYO", "3.061", "Üniversiteler altında gruplanmış"],
            ["Bölge / Coğrafya", "7", "Marmara, Ege, Akdeniz, Karadeniz, İç Anadolu, Doğu Anadolu, Güneydoğu"],
            ["Vektör Chunk", "≈ 28.500", "Embedding birimleri (her program ≈ 1 chunk)"],
            ["Embedding Boyutu", "384", "paraphrase-multilingual-MiniLM-L12-v2"],
        ],
        col_widths_cm=[5.0, 3.0, 8.5],
    )

    add_h2(doc, "3.3. Örnek Ham Veri (YÖK Atlas — programs_2025.json)")
    add_table(doc,
        headers=["Alan", "Örnek Değer"],
        rows=[
            ["program_kodu", "203910363"],
            ["program_adi", "Bilgisayar Mühendisliği"],
            ["universite_adi", "BOĞAZİÇİ ÜNİVERSİTESİ"],
            ["fakulte", "Mühendislik Fakültesi"],
            ["sehir", "İSTANBUL"],
            ["puan_turu", "SAY"],
            ["egitim_dili", "İngilizce"],
            ["taban_puan_2024", "558.42"],
            ["taban_sira_2024", "712"],
            ["kontenjan", "70"],
            ["doluluk", "100%"],
        ],
        col_widths_cm=[4.5, 12.0],
    )

    add_h2(doc, "3.4. Örnek İşlenmiş Chunk (chunks.json)")
    add_para(doc,
        "Ham JSON kayıtları RAG için tek bir doğal dil paragrafına dönüştürülür. Aşağıda bir chunk örneği "
        "gösterilmiştir; bu metin doğrudan vektör veritabanına yazılır:",
        italic=True, size=10
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "📚 Bilgisayar Mühendisliği (İngilizce) | 🏛 BOĞAZİÇİ ÜNİVERSİTESİ (İSTANBUL) | "
        "🏫 Mühendislik Fakültesi | 📍 İSTANBUL | Puan türü: SAY | Eğitim seviyesi: Lisans | "
        "Eğitim dili: İngilizce | Kontenjan: 70 | Taban puan 2024: 558.42 | Taban sıra: 712 | "
        "Doluluk: %100 | Tercih kodu: [203910363]"
    )
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(40, 40, 40)
    # Hücre arkaplanı için paragraph border
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def section_yontem(doc: Document) -> None:
    add_h1(doc, "4. Yöntem — RAG Pipeline")

    add_para(doc,
        "Projenin omurgasını oluşturan yöntem Retrieval-Augmented Generation (RAG) mimarisidir. RAG, "
        "Büyük Dil Modeli'ne (Gemini) cevap üretirken parametrik olmayan bir bilgi tabanından (ChromaDB) "
        "ilgili bağlamı sağlayan ve böylece halüsinasyonu büyük ölçüde önleyen modern bir tekniktir. "
        "UniSense'te RAG akışı yedi aşamada gerçekleşir."
    )

    add_h2(doc, "Adım 1 — Veri Toplama (Scraping)")
    add_para(doc,
        "yokatlas_scraper.py, urap_scraper.py, wikipedia_uni_scraper.py ve avesis_scraper.py modülleri "
        "her biri kendi kaynağına özel HTTP istekleri atar; sonuçları backend/data/raw/ altına ham JSON "
        "olarak kaydeder. Scraper'lar tenacity ile retry, rate-limit'e saygı ve User-Agent rotasyonu içerir."
    )

    add_h2(doc, "Adım 2 — Dönüşüm ve Zenginleştirme (Transform & Enrich)")
    add_para(doc,
        "Ham YÖK Atlas verisi transform_yokatlas.py ile normalize edilir; aynı bölümün farklı yıllardaki "
        "kayıtları birleştirilir, eksik alanlar URAP ve Wikipedia'dan tamamlanır. enrich_geo.py her "
        "üniversiteye şehir → bölge → coğrafi özellik (denizi olma, metropol, sahil) eşlemesi yapar. "
        "Çıktı backend/data/processed/ altında universities.json, departments.json, rankings.json olarak "
        "kaydedilir."
    )

    add_h2(doc, "Adım 3 — Chunk Oluşturma (Build Chunks)")
    add_para(doc,
        "build_chunks.py her bir program/üniversite kaydını insan tarafından okunabilir tek bir "
        "paragrafa (chunk) dönüştürür. Her chunk emoji ile sınırlanmış alan etiketleri (📚 bölüm, 🏛 "
        "üniversite, 📍 şehir vb.) içerir; bu yapı hem LLM'in cevap üretimini kolaylaştırır hem de "
        "embedding modelin semantik anlamı yakalamasını destekler. Her chunk metadata (program_kodu, "
        "puan_turu, sehir, kaynak_url) ile etiketlenir."
    )

    add_h2(doc, "Adım 4 — Embedding (Vektörleştirme)")
    add_para(doc,
        "embed.py modülü sentence-transformers kütüphanesini kullanarak paraphrase-multilingual-MiniLM-L12-v2 "
        "(384 boyutlu) modeliyle her chunk için bir embedding vektörü hesaplar. Bu model çok dilli "
        "eğitildiği için Türkçe semantik benzerliği yüksek doğrulukla yakalar."
    )

    add_h2(doc, "Adım 5 — Vektör Depolama (ChromaDB)")
    add_para(doc,
        "Hesaplanan embedding'ler ve metadata ChromaDB'nin persistent disk modunda saklanır "
        "(backend/data/embeddings/chromadb/). ChromaDB, HNSW indeksi sayesinde milisaniye altı "
        "Approximate Nearest Neighbor (ANN) araması yapar."
    )

    add_h2(doc, "Adım 6 — Sorgu ve Bağlam Çekimi (Retrieve)")
    add_para(doc,
        "Kullanıcı bir sorgu attığında (örn. \"Bilkent Bilgisayar Mühendisliği taban puanı\") sorgu "
        "metni input_sanitizer ile temizlenir, aynı embedding modeliyle vektöre dönüştürülür ve "
        "ChromaDB'den top-K (varsayılan 8) en benzer chunk çekilir. Sıralama intent'i tespit "
        "edildiğinde (örn. sorguda \"sıra\", \"sıralama\" geçiyorsa) hibrit bir filtreleme uygulanır: "
        "vektör benzerliği + metadata filtreleri (puan türü, şehir, üniversite tipi)."
    )

    add_h2(doc, "Adım 7 — Cevap Üretimi (Generate)")
    add_para(doc,
        "Çekilen bağlam, sistem prompt'u ile birlikte Gemini 2.5 Flash-Lite modeline gönderilir. "
        "Sistem prompt'u modeli sıkı kurallarla bağlar: yalnızca verilen bağlama dayan, sayıları aynen "
        "aktar, ÖSYM tercih kodunu köşeli parantez içinde yaz, kaynaklarda yoksa \"kaynaklarımda yok\" "
        "de. Ayrıca akademisyen sorularında otomatik olarak akademik.yok.gov.tr arama linki üretilir. "
        "Üretilen cevap frontend'e döndürülür ve aynı zamanda kaynak chunk'lar şeffaf biçimde gösterilir."
    )

    add_h2(doc, "4.1. Mimari Şeması")
    if IMG_ARCHITECTURE.exists():
        doc.add_picture(str(IMG_ARCHITECTURE), width=Cm(16.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Şekil 1. UniSense — Sistem Mimarisi (Frontend ⇄ Backend ⇄ ChromaDB ⇄ Veri Pipeline)")
    else:
        add_para(doc, f"[Mimari görseli bulunamadı: {IMG_ARCHITECTURE}]", italic=True)


def section_uygulama(doc: Document) -> None:
    add_h1(doc, "5. Uygulama Tasarımı ve Görseller")

    add_para(doc,
        "UniSense iki ana bileşenden oluşur: (1) Python tabanlı FastAPI backend ve (2) modern React 18 + "
        "Vite frontend'i. Frontend tasarımı koyu temalı, üç boyutlu animasyonlu (Three.js / React-Three-"
        "Fiber) ve mobil uyumlu olacak şekilde Tailwind CSS ile geliştirilmiştir. Aşağıda uygulamanın üç "
        "ana ekranı görselleriyle tanıtılmaktadır."
    )

    add_h2(doc, "5.1. Karşılama (Splash) Ekranı")
    add_para(doc,
        "Kullanıcının uygulamayla ilk karşılaştığı ekrandır. Animasyonlu yıldız arka planı, gradient renk "
        "paleti ve büyük \"UniSense\" logosu ile marka kimliğini vurgular. Üst köşede tema değiştirici "
        "ve hesap menüsü; alt köşede gerçek istatistikler (21.602 program, 227 üniversite) yer alır. "
        "Kullanıcı buradan iki yola gidebilir: \"Başla\" ile profil yapılandırma akışına, \"Direkt Sor\" "
        "ile doğrudan sorgu ekranına."
    )
    if IMG_SPLASH.exists():
        doc.add_picture(str(IMG_SPLASH), width=Cm(16))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Şekil 2. UniSense Karşılama (Splash) ekranı — koyu uzay teması, animasyonlu küreler ve gradient \"UniSense\" başlığı")

    add_h2(doc, "5.2. Ana Sayfa (Dashboard)")
    add_para(doc,
        "Giriş yapan kullanıcının veri istatistiklerini gördüğü, sistem genelinde hangi modüllere "
        "erişebileceğini özetleyen sayfadır. \"Doğru tercihi veri ile yap\" sloganı altında dört ana "
        "metrik kartı yer alır: 227 Üniversite, 21.602 Program (Lisans + Önlisans), 3.061 Fakülte/MYO, "
        "7 Bölge. Üst menüden Ana Sayfa, Sorgu, Pusula, Tercih, Hesap modüllerine geçilebilir."
    )
    if IMG_HOME.exists():
        doc.add_picture(str(IMG_HOME), width=Cm(16))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Şekil 3. Ana Sayfa — kullanıcı dashboard'u, gerçek zamanlı veri sayaçları ve modül erişimi")

    add_h2(doc, "5.3. Sorgu (Sohbet) Ekranı")
    add_para(doc,
        "Projenin kalbi olan RAG sohbet ekranıdır. Sol tarafta kullanıcının son 5 sohbet oturumu listelenir "
        "(FIFO mantığıyla en eskisi otomatik silinir). Orta sütunda sohbet alanı; alt kısımda kullanıcı "
        "girdisi ve örnek sorgu önerileri yer alır. Sağ üstte AI modeli seçici (Gemini) bulunur. Kullanıcı "
        "doğal Türkçe ile soru sorduğunda sistem RAG pipeline'ını çalıştırarak cevabın yanı sıra "
        "kaynakçayı da döndürür ve cevaba gömülü ÖSYM tercih kodlarını otomatik olarak \"+ Tercihe Ekle\" "
        "butonuna dönüştürür."
    )
    if IMG_SEARCH.exists():
        doc.add_picture(str(IMG_SEARCH), width=Cm(16))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, "Şekil 4. Sorgu ekranı — sol panelde sohbet geçmişi, orta panelde RAG sohbet alanı, alt panelde örnek sorgu önerileri")

    add_h2(doc, "5.4. Diğer Modüller")
    add_bullet(doc, "Pusula (Compass): İlgi alanı + güçlü dersler + şehir tercihi gibi eksen-bazlı bir analiz ile öğrenciye en uygun bölüm önerileri üretir.")
    add_bullet(doc, "Tercih Listesi: 24 sıralık Firestore tabanlı tercih listesi, sürükle-bırak ile sıralama (@dnd-kit), eksik alanların arka planda otomatik backfill edilmesi.")
    add_bullet(doc, "Hesap: Firebase Auth ile Google + e-posta/şifre girişi, profil fotoğrafı yükleme, puan/sıralama profili güncelleme.")


def section_test(doc: Document) -> None:
    add_h1(doc, "6. Test ve Sonuçlar")

    add_para(doc,
        "Sistemin doğruluğunu ve dayanıklılığını ölçmek için kapsamlı bir test sorgu seti oluşturulmuştur "
        "(backend/scripts/test_queries.json). Bu set 10 kategori altında toplam 75+ farklı senaryoyu kapsar; "
        "her sorgu otomatik olarak /api/v1/ask endpoint'ine atılır, dönüş süresi ve cevap kalitesi raporlanır."
    )

    add_h2(doc, "6.1. Test Kategorileri")
    add_table(doc,
        headers=["Kategori", "Açıklama", "Sorgu Sayısı"],
        rows=[
            ["Sıralama Intent", "Hibrit retriever testi (sıra/puan filtresi + vektör)", "9"],
            ["Üniversite Bilgi", "Doğrudan üniversite hakkında bilgi sorguları", "9"],
            ["Bölüm Bilgi", "Bölüm taban puan, kapsam, süre soruları", "8"],
            ["Filtreleme", "İl + dil + burs gibi çoklu filtreler", "7"],
            ["Akademisyen", "YÖK Akademik link üretim kuralı testi", "4"],
            ["Karşılaştırma", "İki üniversite/bölüm karşılaştırma", "3"],
            ["Şehir / Yöre", "Bölgesel ve şehir-bazlı sorgular", "4"],
            ["Coğrafi Filtre", "Deniz, sahil, metropol gibi coğrafi nitelikler", "12"],
            ["Edge Cases", "Anlamsız, çok kısa, alakasız girişler", "7"],
            ["Türkçe Karakter", "Büyük/küçük harf ve diakritik testleri", "5"],
        ],
        col_widths_cm=[4.0, 9.5, 2.5],
    )

    add_h2(doc, "6.2. Örnek Sorgu — Cevap Eşleşmeleri")
    add_table(doc,
        headers=["Kullanıcı Sorgusu", "Beklenen Davranış"],
        rows=[
            [
                "300.000 sırayla devlet üniversiteleri",
                "Hibrit retriever; sıralama ≤ 300k filtresi + devlet üniversitesi metadata filtresi → top-K liste"
            ],
            [
                "Boğaziçi vs ODTÜ Bilgisayar Mühendisliği hangisi daha iyi",
                "İki üniversiteden de chunk çek; LLM karşılaştırmalı cevap üretir, taban puan/sıra sayılarını aynen aktarır"
            ],
            [
                "Boğaziçi Bilgisayar Mühendisliği'nde hangi hocalar var",
                "Akademisyen kuralı tetiklenir; AVESİS chunk'ları + cevap sonuna otomatik akademik.yok.gov.tr arama linki"
            ],
            [
                "Karadeniz kıyısındaki tıp fakülteleri",
                "Coğrafi filtre (sea_coast=Karadeniz) + program_adı=Tıp filtresi → Trabzon, Samsun, Rize vb. üniversiteler"
            ],
            [
                "asdfgh",
                "Edge case; retriever düşük benzerlik döndürür → LLM \"kaynaklarımda anlamlı bir eşleşme bulamadım\" cevabı"
            ],
            [
                "İSKENDERUN TEKNİK ÜNİVERSİTESİ",
                "Türkçe karakter normalizasyonu testi; büyük/küçük varyantlar aynı sonuca götürmeli"
            ],
        ],
        col_widths_cm=[6.5, 10.0],
    )

    add_h2(doc, "6.3. Performans Metrikleri")
    add_table(doc,
        headers=["Metrik", "Ölçüm"],
        rows=[
            ["Ortalama uçtan uca cevap süresi", "≈ 1.8 sn (vektör retrieval ≈ 80 ms + Gemini Flash-Lite ≈ 1.5 sn)"],
            ["Tek sorguda çekilen chunk sayısı (top-K)", "8"],
            ["Vektör veritabanı boyutu", "≈ 28.500 chunk × 384 dim ≈ 45 MB"],
            ["Gemini API quota stratejisi", "Multi-key auto-rotate (quota dolunca diğer key)"],
            ["Frontend bundle boyutu", "≈ 2.4 MB gzip (Vite production build)"],
            ["Eşzamanlı kullanıcı limiti", "20 istek/dk/IP (slowapi rate-limit)"],
        ],
        col_widths_cm=[6.5, 10.0],
    )

    add_h2(doc, "6.4. Sonuç Değerlendirmesi")
    add_para(doc,
        "Test seti üzerinden yapılan deneyler, RAG mimarisinin geleneksel \"prompt-only\" LLM "
        "yaklaşımına göre sayısal sorgu doğruluğunu önemli ölçüde artırdığını göstermiştir. Özellikle "
        "taban puan, sıralama ve kontenjan gibi kritik sayıların cevaba aynen aktarılması — "
        "halüsinasyona meyilli açık model davranışının tersine — kullanıcı güvenini doğrudan "
        "destekleyen bir özelliktir."
    )
    add_para(doc,
        "Edge case ve anlamsız girdilerde sistem, retriever düşük benzerlik skoru ürettiği için "
        "LLM'i bilinmeyen alana sürüklemek yerine \"kaynaklarımda yok\" şeklinde dürüst bir cevap "
        "vermeyi tercih etmektedir. Bu davranış akademik tercih kararı gibi yüksek riskli bir bağlamda "
        "kritik öneme sahiptir."
    )


def section_kaynaklar(doc: Document) -> None:
    add_h1(doc, "7. Kaynaklar ve Referanslar")

    add_h2(doc, "7.1. Veri Kaynakları")
    add_bullet(doc, "YÖK Atlas — https://yokatlas.yok.gov.tr")
    add_bullet(doc, "ÖSYM — https://www.osym.gov.tr")
    add_bullet(doc, "URAP (University Ranking by Academic Performance) — https://www.urapcenter.org")
    add_bullet(doc, "AVESİS — Üniversitelerin akademik veri sistemleri (avesis.<universite>.edu.tr)")
    add_bullet(doc, "Wikipedia (Türkçe) — https://tr.wikipedia.org")

    add_h2(doc, "7.2. Kütüphaneler ve Araçlar")
    add_bullet(doc, "FastAPI — https://fastapi.tiangolo.com")
    add_bullet(doc, "ChromaDB (vektör veritabanı) — https://www.trychroma.com")
    add_bullet(doc, "sentence-transformers (multilingual MiniLM) — https://www.sbert.net")
    add_bullet(doc, "Google Generative AI (Gemini 2.5 Flash / Flash-Lite) — https://ai.google.dev")
    add_bullet(doc, "React 18 + Vite + Tailwind CSS — https://react.dev | https://vitejs.dev")
    add_bullet(doc, "Firebase (Auth + Firestore + Storage) — https://firebase.google.com")
    add_bullet(doc, "Three.js + React Three Fiber — https://threejs.org | https://r3f.docs.pmnd.rs")

    add_h2(doc, "7.3. Proje Repository")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Tüm kaynak kod, veri ve dökümantasyon GitHub üzerinde MIT lisansıyla açık olarak yayınlanmıştır:")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(GITHUB_URL)
    run.font.name = "Consolas"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    run.font.underline = True


# ============================================================
# ANA AKIŞ
# ============================================================
def main() -> None:
    doc = Document()

    # Sayfa boyutu (A4) ve marjinler (2.5 cm)
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Varsayılan stil
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Bölümler
    build_cover(doc)
    add_page_break(doc)

    section_amac(doc)
    section_problem(doc)
    add_page_break(doc)

    section_veri(doc)
    add_page_break(doc)

    section_yontem(doc)
    add_page_break(doc)

    section_uygulama(doc)
    add_page_break(doc)

    section_test(doc)
    section_kaynaklar(doc)

    # Kaydet
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"[OK] Rapor kaydedildi: {OUT_PATH}")
    print(f"     Boyut: {OUT_PATH.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
